#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""규격 md -> 검토 전달용 docx 변환기.

    python tools/md_to_docx.py <입력.md> <출력.docx>

**왜 있나.** 운영자 검토는 docx 왕복으로 돈다(`__review/` — 전달본을 주면
`*_revision.docx` 로 돌아온다). 개정판마다 전달본을 만들어야 하는데(운영자
지시 2026-08-21), pandoc 없는 환경에서도 돌도록 python-docx 로 직접 만든다.

**다루는 markdown.** 이 저장소 규격 문서가 실제로 쓰는 부분집합만이다 —
제목(#~####), 표(| 셀 |, 셀 안 `<br>` 줄바꿈, 백틱 안 `|` 보호), 인용(>),
목록(-, 1.), 코드 펜스, 인라인 **굵게**/`코드`/~~취소~~/*기울임*.
그 밖의 문법은 평문으로 떨어진다 — 조용히 깨지는 것보다 낫다.
"""

from __future__ import annotations

import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

BODY_FONT = 'Malgun Gothic'
CODE_FONT = 'Consolas'


def _set_fonts(run, code=False):
    f = run.font
    f.name = CODE_FONT if code else BODY_FONT
    # python-docx 는 ascii/hAnsi 만 만진다 -- 한글은 eastAsia 폰트를 봐야 한다
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), BODY_FONT)


# ---------------------------------------------------------------- inline
# 코드 스팬을 먼저 갈라야 `**` 같은 것이 코드 안에서 서식으로 오해되지 않는다.

_TOKEN = re.compile(r'(\*\*.+?\*\*|~~.+?~~|\*(?!\s)[^*]+?(?<!\s)\*)', re.S)


def _emit(par, text, *, bold=False, strike=False, italic=False, code=False,
          size=None):
    if not text:
        return
    run = par.add_run(text)
    run.bold = bold or None
    run.italic = italic or None
    run.font.strike = strike or None
    _set_fonts(run, code=code)
    if size:
        run.font.size = size


def _rich(par, text, size=None):
    """인라인 서식을 run 으로 푼다."""
    parts = text.split('`')
    for i, part in enumerate(parts):
        if i % 2 == 1:                      # 백틱 안 -- 코드 스팬
            _emit(par, part, code=True, size=size)
            continue
        pos = 0
        for m in _TOKEN.finditer(part):
            _emit(par, part[pos:m.start()], size=size)
            tok = m.group(0)
            if tok.startswith('**'):
                _rich_nested(par, tok[2:-2], bold=True, size=size)
            elif tok.startswith('~~'):
                _emit(par, tok[2:-2], strike=True, size=size)
            else:
                _emit(par, tok[1:-1], italic=True, size=size)
            pos = m.end()
        _emit(par, part[pos:], size=size)


def _rich_nested(par, text, *, bold=False, size=None):
    # 굵은 구간 안의 `코드` 만 한 겹 더 처리한다
    parts = text.split('`')
    for i, part in enumerate(parts):
        _emit(par, part, bold=bold, code=(i % 2 == 1), size=size)


# ---------------------------------------------------------------- 표

def _split_row(line):
    """`|` 로 셀을 가른다 -- 백틱 코드 스팬 안의 `|` 는 보호한다."""
    cells, buf, in_code = [], [], False
    for ch in line.strip().strip('|') + '|':
        if ch == '`':
            in_code = not in_code
            buf.append(ch)
        elif ch == '|' and not in_code:
            cells.append(''.join(buf).strip())
            buf = []
        else:
            buf.append(ch)
    return cells


_SEP = re.compile(r'^:?-{3,}:?$')


def _is_sep_row(cells):
    return cells and all(_SEP.match(c) for c in cells)


def _fill_cell(cell, text, *, header=False):
    cell.paragraphs[0].text = ''
    first = True
    for chunk in text.split('<br>'):
        par = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        par.paragraph_format.space_before = Pt(1)
        par.paragraph_format.space_after = Pt(1)
        if header:
            _rich_nested(par, chunk.strip(), bold=True, size=Pt(9))
        else:
            _rich(par, chunk.strip(), size=Pt(9))


def _add_table(doc, rows):
    parsed = [_split_row(r) for r in rows]
    body = [c for c in parsed if not _is_sep_row(c)]
    if not body:
        return
    ncol = len(body[0])
    table = doc.add_table(rows=0, cols=ncol)
    table.style = 'Table Grid'
    table.autofit = True
    for ri, cells in enumerate(body):
        row = table.add_row()
        for ci in range(ncol):
            _fill_cell(row.cells[ci], cells[ci] if ci < len(cells) else '',
                       header=(ri == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------- 본문

def convert(src, dst):
    with open(src, encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
    for s in doc.sections:
        s.page_width, s.page_height = Mm(210), Mm(297)
        s.left_margin = s.right_margin = Mm(18)
        s.top_margin = s.bottom_margin = Mm(18)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        m = re.match(r'^(#{1,4})\s+(.*)$', line)
        if m:
            par = doc.add_heading('', level=len(m.group(1)))
            _rich_nested(par, m.group(2))
            for run in par.runs:
                _set_fonts(run)
            i += 1
            continue

        if line.startswith('```'):
            i += 1
            buf = []
            while i < n and not lines[i].startswith('```'):
                buf.append(lines[i])
                i += 1
            i += 1
            for code_line in buf:
                par = doc.add_paragraph()
                par.paragraph_format.space_after = Pt(0)
                par.paragraph_format.left_indent = Mm(6)
                _emit(par, code_line or ' ', code=True, size=Pt(8.5))
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        if line.lstrip().startswith('|'):
            rows = []
            while i < n and lines[i].lstrip().startswith('|'):
                rows.append(lines[i])
                i += 1
            _add_table(doc, rows)
            continue

        if line.startswith('>'):
            while i < n and lines[i].startswith('>'):
                inner = lines[i][1:].lstrip()
                i += 1
                if not inner:
                    continue
                par = doc.add_paragraph()
                par.paragraph_format.left_indent = Mm(8)
                par.paragraph_format.space_after = Pt(2)
                pf = par.paragraph_format
                pf.space_before = Pt(0)
                for run_text in (inner,):
                    _rich(par, run_text, size=Pt(9.5))
                for run in par.runs:
                    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            continue

        m = re.match(r'^\s*[-*]\s+(.*)$', line)
        if m:
            par = doc.add_paragraph(style='List Bullet')
            _rich(par, m.group(1))
            i += 1
            continue

        m = re.match(r'^\s*(\d+)[.)]\s+(.*)$', line)
        if m:
            par = doc.add_paragraph(style='List Number')
            _rich(par, m.group(2))
            i += 1
            continue

        par = doc.add_paragraph()
        _rich(par, line.strip())
        i += 1

    doc.save(dst)
    print(f'wrote {dst}')


if __name__ == '__main__':
    if len(sys.argv) != 3:
        sys.exit(__doc__)
    convert(sys.argv[1], sys.argv[2])
